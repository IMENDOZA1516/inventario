from flask import Flask, request, jsonify, url_for, session, render_template
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from extensions import db
from flask_jwt_extended import JWTManager, create_access_token
from models import Computadora, Empleado, Usuario
from flask_migrate import Migrate  # Importar Flask-Migrate
from flask import Flask, render_template
from flask import session, redirect, url_for, render_template
from flask import Response
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
import pandas as pd
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from flask import send_file
from models import db, Reasignacion  # Importamos la base de datos y el modelo
from flask import Flask, render_template, request, redirect, url_for, flash
from datetime import datetime
import os
from docx import Document
from docx.shared import Pt
from io import BytesIO
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from models import db, Computadora, ComponenteDefectuoso
from flask_login import login_required, current_user
from docx.oxml.ns import qn
from flask_login import LoginManager
from models import ComputadoraEliminada
from flask import Flask, request, jsonify, render_template, redirect, url_for, session
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from models import Computadora, Accesorio
from sqlalchemy import text
from models import Migracion  # Si los modelos están en un archivo llamado models.py
from werkzeug.security import generate_password_hash
from models import Usuario  # ajusta si tu modelo está en otro archivo


app = Flask(__name__)


# Configuración de JWT
app.config['JWT_SECRET_KEY'] = 'Mendoza0101@'  # Cambia esto por una clave segura
jwt = JWTManager(app)

app.secret_key = 'Mendoza0101'  # Usa una clave única y segura

# Leer directamente de la variable de entorno sin poner un valor por defecto
# Configuración de la base de datos desde variable de entorno o valor por defecto local
app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv("DATABASE_URL", "postgresql://postgres:viqVSHTJagoHbtYkytwtKFapwPsSqiEi@shinkansen.proxy.rlwy.net:47285/railway")
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = os.getenv("SECRET_KEY", "Mendoza0101")
db.init_app(app)

# Configurar Flask-Migrate
migrate = Migrate(app, db)
# 🔥 INICIALIZAR FLASK-LOGIN 🔥
login_manager = LoginManager()
login_manager.init_app(app)  # Asociar Flask-Login con la app
login_manager.login_view = "login"  # Si no está autenticado, redirigir al login

# 🔥 FUNCIÓN PARA CARGAR USUARIOS 🔥
@login_manager.user_loader
def load_user(user_id):
    return Usuario.query.get(int(user_id))  # Busca usuarios por ID en la base de datos

@app.route('/')
def index():
    return redirect(url_for('login'))

# ---------- Ruta para agregar computadora y crear empleado + usuario ----------
# ---------- Ruta para agregar computadora asignada a un empleado existente ----------
@app.route('/computadoras', methods=['POST'])
def agregar_computadora():
    data = request.get_json()

    empleado_id = data.get('empleado_id')
    if not empleado_id:
        return jsonify({'mensaje': 'Error: Debes seleccionar un empleado'}), 400

    # Crear la computadora
    nueva_computadora = Computadora(
        numero_inventario=data.get('numero_inventario'),
        ubicacion_campus=data.get('ubicacion_campus'),
        departamento=data.get('departamento'),
        monitor=data.get('monitor'),
        teclado=data.get('teclado'),
        estado=data.get('estado'),
        tipo_propiedad=data.get('tipo_propiedad'),
        empleado_id=empleado_id,
        monitor_obsoleto=data.get('monitor_obsoleto', False),
        teclado_obsoleto=data.get('teclado_obsoleto', False)
    )
    db.session.add(nueva_computadora)
    db.session.commit()  # 👈 Guardamos para que tenga ID antes de asignar accesorios

    # Accesorios (si vienen)
    accesorios = data.get('accesorios', [])
    for acc in accesorios:
        tipo = acc.get('tipo')
        inventario = acc.get('numero_inventario')
        ligado_a_pc = acc.get('ligado_a_pc', True)  # por defecto, se liga a la computadora

        if tipo and inventario:
            accesorio = Accesorio(
                tipo=tipo,
                numero_inventario=inventario,
                empleado_id=empleado_id,
                computadora_id=nueva_computadora.id if ligado_a_pc else None
            )
            db.session.add(accesorio)

    db.session.commit()

    return jsonify({'mensaje': 'Computadora y accesorios agregados correctamente'}), 201


#---- ruta para agregar empleado + usuario----
@app.route('/empleados', methods=['POST'])
def agregar_empleado():
    data = request.get_json()

    # Validar existencia por correo
    empleado_existente = Empleado.query.filter_by(email=data.get('email')).first()
    usuario_existente = Usuario.query.filter_by(email=data.get('email')).first()
    if empleado_existente or usuario_existente:
        return jsonify({'mensaje': 'Error: El correo ya está registrado'}), 400

    # Si el rol es "usuario", se crea como empleado
    if data.get('rol') == 'usuario':
        nuevo_empleado = Empleado(
            nombre=data.get('nombre'),
            email=data.get('email'),
            puesto=data.get('puesto'),
            campus=data.get('campus')
        )
        db.session.add(nuevo_empleado)
        db.session.commit()

    # Crear usuario (ya sea admin o usuario)
    rol = data.get('rol', 'usuario')  # Por si no viene, usa 'usuario' por defecto

    nuevo_empleado = Empleado(
    nombre=data.get('nombre'),
    email=data.get('email'),
    puesto=data.get('puesto'),
    campus=data.get('campus')
)
    db.session.add(nuevo_empleado)
    db.session.commit()

    nuevo_usuario = Usuario(
    nombre=data.get('nombre'),
    email=data.get('email'),
    rol=rol  # 👈 ahora sí respeta el rol recibido
)
    nuevo_usuario.set_password(data.get('password'))
    db.session.add(nuevo_usuario)
    db.session.commit()

    return jsonify({'mensaje': 'Usuario creado correctamente'}), 201


# ---------- Ruta para obtener empleados ----------
@app.route('/empleados', methods=['GET'])
def obtener_empleados():
    empleados = Empleado.query.all()
    empleados_json = [{
        'id': emp.id,
        'nombre': emp.nombre,
        'email': emp.email,
        'puesto': emp.puesto,
        'campus': emp.campus
    } for emp in empleados]
    return jsonify(empleados_json)

# ---------- Buscar computadoras ----------
@app.route('/buscar_computadoras')
def buscar_computadoras():
    query = request.args.get('q', '').strip().lower()
    tipo_propiedad = request.args.get('tipo', '').strip().capitalize()  # Normalizamos "propia" y "alquilada"

    computadoras_query = Computadora.query.filter(Computadora.es_obsoleta.is_(False))  # 🔹 Evita filter_by

    if query:
        computadoras_query = computadoras_query.filter(
            (Computadora.numero_inventario.ilike(f"%{query}%")) |
            (Computadora.empleado.has(Empleado.nombre.ilike(f"%{query}%")))
        )

    if tipo_propiedad in ["Propia", "Alquilada"]:  # 🔹 Ahora "Alquilada" es correcto
        computadoras_query = computadoras_query.filter(Computadora.tipo_propiedad.ilike(tipo_propiedad))

    computadoras = computadoras_query.all()
    computadoras_filtradas = [{
        'id': comp.id,
        'numero_inventario': comp.numero_inventario,
        'estado': comp.estado,
        'departamento': comp.departamento,
        'ubicacion_campus': comp.ubicacion_campus,
        'monitor': comp.monitor,
        'teclado': comp.teclado,
        'tipo_propiedad': comp.tipo_propiedad,
        'empleado': {
            'nombre': comp.empleado.nombre,
            'email': comp.empleado.email,
            'puesto': comp.empleado.puesto
        } if comp.empleado else None
    } for comp in computadoras]

    return jsonify(computadoras_filtradas)



@app.route('/computadoras', methods=['GET'])
@login_required
def obtener_computadoras():
    page = request.args.get('page', 1, type=int)
    per_page = 10

    rol_usuario = current_user.rol
    if rol_usuario == "usuario":
        empleado = Empleado.query.filter_by(email=current_user.email).first()
        if not empleado:
            return jsonify([])
        query = Computadora.query.filter_by(empleado_id=empleado.id, es_obsoleta=False)
    else:
        query = Computadora.query.filter(Computadora.es_obsoleta.is_(False))

    paginated = query.paginate(page=page, per_page=per_page, error_out=False)

    lista_computadoras = []
    for comp in paginated.items:
        computadora_info = {
            'id': comp.id,
            'numero_inventario': comp.numero_inventario,
            'estado': comp.estado,
            'departamento': comp.departamento,
            'ubicacion_campus': comp.ubicacion_campus,
            'monitor': comp.monitor,
            'teclado': comp.teclado,
            'tipo_propiedad': comp.tipo_propiedad,
            'empleado_id': comp.empleado_id,
            'empleado': {
                'id': comp.empleado.id,
                'nombre': comp.empleado.nombre,
                'email': comp.empleado.email,
                'puesto': comp.empleado.puesto,
                'campus': comp.empleado.campus
            } if comp.empleado else None,
            'accesorios': [
                {
                    'tipo': a.tipo,
                    'numero_inventario': a.numero_inventario
                } for a in comp.accesorios
            ],
            # 👇 Agregamos los componentes obsoletos
            'componentes_defectuosos': [
                {
                    'tipo': d.tipo_componente,
                    'inventario': d.inventario_componente
                } for d in comp.componentes_defectuosos
            ]
        }
        lista_computadoras.append(computadora_info)

    return jsonify({
        'computadoras': lista_computadoras,
        'total_paginas': paginated.pages,
        'pagina_actual': paginated.page
    })


@app.route('/computadoras/<int:id>', methods=['PUT'])
def actualizar_computadora(id):
    computadora = Computadora.query.get(id)
    if not computadora:
        return jsonify({'error': 'Computadora no encontrada'}), 404

    data = request.json

    computadora.numero_inventario = data.get('numero_inventario', computadora.numero_inventario)
    computadora.ubicacion_campus = data.get('ubicacion_campus', computadora.ubicacion_campus)
    computadora.departamento = data.get('departamento', computadora.departamento)
    computadora.tipo_propiedad = data.get('tipo_propiedad', computadora.tipo_propiedad)
    computadora.estado = data.get('estado', computadora.estado)

    empleado_id = data.get('empleado_id')
    if empleado_id:
        computadora.empleado_id = empleado_id

    # Monitor obsoleto
    if data.get('monitor_obsoleto') and computadora.monitor:
        comp_def = ComponenteDefectuoso(
            tipo_componente="Monitor",
            modelo=computadora.monitor,
            inventario_componente=computadora.monitor,
            motivo=data.get('motivo_monitor') or "Daño irreparable",
            computadora_id=computadora.id
        )
        db.session.add(comp_def)
        computadora.monitor = None
        computadora.monitor_obsoleto = True
    else:
        computadora.monitor = data.get('monitor')
        computadora.monitor_obsoleto = False

    # Teclado obsoleto
    if data.get('teclado_obsoleto') and computadora.teclado:
        comp_def = ComponenteDefectuoso(
            tipo_componente="Teclado",
            modelo=computadora.teclado,
            inventario_componente=computadora.teclado,
            motivo=data.get('motivo_teclado') or "Daño irreparable",
            computadora_id=computadora.id
        )
        db.session.add(comp_def)
        computadora.teclado = None
        computadora.teclado_obsoleto = True
    else:
        computadora.teclado = data.get('teclado')
        computadora.teclado_obsoleto = False

    # CPU obsoleta
    if data.get('cpu_obsoleta'):
        existe = ComponenteDefectuoso.query.filter_by(
            computadora_id=computadora.id,
            tipo_componente="CPU",
            inventario_componente=computadora.numero_inventario
        ).first()
        if not existe:
            cpu_def = ComponenteDefectuoso(
                tipo_componente="CPU",
                modelo="N/A",
                inventario_componente=computadora.numero_inventario,
                motivo=data.get('motivo_cpu') or "Daño irreparable",
                computadora_id=computadora.id
            )
            db.session.add(cpu_def)

    # ✅ SOLO eliminamos accesorios ligados a esta computadora
    Accesorio.query.filter_by(computadora_id=computadora.id).delete()

    # ✅ Insertamos los nuevos accesorios
    for acc in data.get('accesorios', []):
        tipo = acc.get('tipo')
        inventario = acc.get('numero_inventario')
        ligado_a_pc = acc.get('ligado_a_pc', True)
        es_obsoleto = acc.get('obsoleto', False)
        motivo = acc.get('motivo') or "Daño irreparable"

        if tipo and inventario:
            if es_obsoleto:
                comp_def = ComponenteDefectuoso(
                    tipo_componente=tipo,
                    modelo=tipo,
                    inventario_componente=inventario,
                    motivo=motivo,
                    computadora_id=computadora.id
                )
                db.session.add(comp_def)
            else:
                accesorio = Accesorio(
                    tipo=tipo,
                    numero_inventario=inventario,
                    empleado_id=computadora.empleado_id,
                    computadora_id=computadora.id if ligado_a_pc else None
                )
                db.session.add(accesorio)

    db.session.commit()
    return jsonify({'mensaje': 'Computadora actualizada correctamente'})





@app.route('/computadoras/<int:id>', methods=['GET'])
@login_required
def obtener_computadora_por_id(id):
    comp = Computadora.query.get(id)
    if not comp:
        return jsonify({'error': 'Computadora no encontrada'}), 404

    # ✅ Obtener accesorios del empleado: ligados y no ligados
    accesorios = []
    if comp.empleado:
        for a in comp.accesorios:  # Solo los accesorios ligados a esta computadora
            accesorios.append({
            'tipo': a.tipo,
            'numero_inventario': a.numero_inventario,
            'ligado_a_pc': True  # Ya sabemos que sí lo están
        })

    comp_json = {
        'id': comp.id,
        'numero_inventario': comp.numero_inventario,
        'ubicacion_campus': comp.ubicacion_campus,
        'departamento': comp.departamento,
        'estado': comp.estado,
        'monitor': comp.monitor,
        'teclado': comp.teclado,
        'tipo_propiedad': comp.tipo_propiedad,
        'monitor_obsoleto': comp.monitor_obsoleto,
        'teclado_obsoleto': comp.teclado_obsoleto,
        'empleado': {
            'nombre': comp.empleado.nombre if comp.empleado else '',
            'email': comp.empleado.email if comp.empleado else ''
        },
        'accesorios': accesorios
    }

    return jsonify(comp_json)




# ---------- Eliminar computadora ----------
@app.route('/computadoras/<int:id>', methods=['DELETE'])
def eliminar_computadora(id):
    computadora = Computadora.query.get(id)
    if not computadora:
        return jsonify({"mensaje": "Computadora no encontrada"}), 404

    # Crear entrada en historial de eliminadas
    eliminada = ComputadoraEliminada(
        numero_inventario=computadora.numero_inventario,
        estado=computadora.estado,
        ubicacion_campus=computadora.ubicacion_campus,
        departamento=computadora.departamento,
        monitor=computadora.monitor,
        teclado=computadora.teclado,
        monitor_obsoleto=computadora.monitor_obsoleto,
        teclado_obsoleto=computadora.teclado_obsoleto,
        tipo_propiedad=computadora.tipo_propiedad,
        empleado_nombre=computadora.empleado.nombre if computadora.empleado else None,
        empleado_email=computadora.empleado.email if computadora.empleado else None
    )

    db.session.add(eliminada)

    # ⚠️ Registrar los accesorios en defectuosos antes de borrar (opcional)
    if computadora.empleado and computadora.empleado.accesorios:
        for acc in computadora.empleado.accesorios:
            componente = ComponenteDefectuoso(
                tipo_componente=acc.tipo,
                modelo=acc.tipo,
                inventario_componente=acc.numero_inventario,
                motivo="Computadora eliminada",
                computadora_id=computadora.id
            )
            db.session.add(componente)

       # ⚠️ Elimina solo los accesorios del empleado que estén ligados a esta computadora
    Accesorio.query.filter_by(
    empleado_id=computadora.empleado.id,
    computadora_id=computadora.id  # 👈 clave: que estén asignados a esta computadora
    ).delete()

    db.session.delete(computadora)
    db.session.commit()

    return jsonify({"mensaje": "Computadora movida a historial de eliminadas"}), 200


# ---------- Marcar computadora como obsoleta ----------
@app.route('/computadoras/<int:id>/obsoleta', methods=['PUT'])
def Obsoletas(id):
    computadora = Computadora.query.get(id)
    if not computadora:
        return jsonify({'error': 'Computadora no encontrada'}), 404

    computadora.es_obsoleta = True
    db.session.commit()
    return jsonify({'mensaje': 'Computadora marcada como obsoleta'}), 200


# ---------- Ruta para obtener computadoras obsoletas ----------
@app.route('/computadoras/obsoletas', methods=['GET'])
def obtener_obsoletas():
    computadoras = Computadora.query.filter_by(es_obsoleta=True).all()
    computadoras_json = [
        {
            'id': comp.id,
            'numero_inventario': comp.numero_inventario,
            'estado': comp.estado,
            'departamento': comp.departamento,
            'ubicacion_campus': comp.ubicacion_campus,
            'monitor': comp.monitor,
            'teclado': comp.teclado
        } for comp in computadoras
    ]
    return jsonify(computadoras_json)


# ---------- Login y logout ----------
@app.route('/login', methods=['POST'])
def login():
    data = request.json
    usuario = Usuario.query.filter_by(email=data['email']).first()

    if usuario and usuario.check_password(data['password']):
        login_user(usuario)
        session['usuario_id'] = usuario.id  # Guarda el ID del usuario en la sesión
        session['usuario'] = usuario.email  # Guarda el email
        session['rol'] = usuario.rol  # Guarda el rol

        return jsonify({"mensaje": "Login exitoso", "redirect": "/accesorios_generales", "rol": usuario.rol}), 200
    else:
        return jsonify({"mensaje": "Credenciales incorrectas"}), 401


@app.route('/login', methods=['GET'])
def login_form():
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    session.clear()
    return redirect(url_for('login'))


# ---------- Obtener rol + ID de empleado ----------
@app.route('/obtener_rol', methods=['GET'])
@login_required
def obtener_rol():
    empleado = Empleado.query.filter_by(email=current_user.email).first()
    empleado_id = empleado.id if empleado else None
    return jsonify({"id": empleado_id, "rol": current_user.rol})


# ---------- Dashboard ----------
@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html', rol=current_user.rol)


# codigo para eleminar la sesion del usuario
#@app.route('/logout')
#@login_required
#def logout():
 #   logout_user()  # 🔥 Cierra sesión correctamente con Flask-Login
  #  session.clear()  # 🔥 Elimina toda la información de sesión
  #  return redirect(url_for('login'))  # Redirigir al login

# rutas para ver las computadoras eliminadas 
@app.route('/eliminadas')
@login_required
def ver_eliminadas():
    return render_template('eliminadas.html')  # este HTML tiene el bloque `{% if request.args.get('fragment') %}`


@app.route('/api/computadoras_eliminadas')
@login_required
def api_computadoras_eliminadas():
    if current_user.rol != 'admin':
        return jsonify({'error': 'Acceso denegado'}), 403

    eliminadas = ComputadoraEliminada.query.all()
    resultado = [
        {
            'numero_inventario': c.numero_inventario,
            'estado': c.estado,
            'departamento': c.departamento,
            'ubicacion_campus': c.ubicacion_campus,
            'monitor': c.monitor,
            'teclado': c.teclado,
            'tipo_propiedad': c.tipo_propiedad,
            'fecha_eliminacion': c.fecha_eliminacion.strftime('%Y-%m-%d %H:%M'),
            'empleado_nombre': c.empleado_nombre or "No asignado"
        }
        for c in eliminadas
    ]
    return jsonify(resultado)





# ---------- Obtener componentes obsoletos ----------
@app.route('/componentes_obsoletos')
@login_required
def ver_componentes_obsoletos():
    return render_template('componentes_obsoletos.html')  # con {% if request.args.get('fragment') %}




@app.route('/api/componentes_obsoletos')
@login_required
def api_componentes_obsoletos():
    if current_user.rol != "admin":
        return jsonify({'error': 'Acceso denegado'}), 403

    # Obtener parámetros de filtro
    inventario = request.args.get('inventario')
    tipo = request.args.get('tipo')
    fecha_desde = request.args.get('fecha_desde')

    # Construir consulta base
    query = ComponenteDefectuoso.query.options(
        db.joinedload(ComponenteDefectuoso.computadora).joinedload(Computadora.empleado),
        db.joinedload(ComponenteDefectuoso.empleado)  # Cargar relación empleado directamente
    )

    if inventario:
        query = query.filter(ComponenteDefectuoso.inventario_componente.ilike(f'%{inventario}%'))

    if tipo:
        query = query.filter_by(tipo_componente=tipo)

    if fecha_desde:
        try:
            fecha_desde = datetime.fromisoformat(fecha_desde)
            query = query.filter(ComponenteDefectuoso.fecha_marcado >= fecha_desde)
        except ValueError:
            pass

    # Ejecutar consulta
    componentes = query.order_by(ComponenteDefectuoso.fecha_marcado.desc()).all()

    componentes_json = []

    for comp in componentes:
        empleado_nombre = None

        # 1. Verificar si tiene empleado directo
        if comp.empleado:
            empleado_nombre = comp.empleado.nombre
        # 2. Intentar obtener desde la computadora
        elif comp.computadora and comp.computadora.empleado:
            empleado_nombre = comp.computadora.empleado.nombre
        # 3. Intentar obtener desde accesorio general (aunque debería estar cubierto por empleado directo)
        
        componentes_json.append({
            'id': comp.id,
            'tipo_componente': comp.tipo_componente,
            'modelo': comp.modelo,
            'inventario': comp.inventario_componente,
            'motivo': comp.motivo,
            'fecha_marcado': comp.fecha_marcado.strftime('%Y-%m-%d %H:%M'),
            'computadora_id': comp.computadora_id,
            'numero_inventario_computadora': comp.computadora.numero_inventario if comp.computadora else "No encontrada",
            'empleado': empleado_nombre or "Sin asignar"
        })

    return jsonify(componentes_json)



@app.route('/reporte/pdf')
def generar_reporte_pdf():
    # Crear un buffer para el PDF
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    pdf.setTitle("Reporte de Computadoras")

    # Título del reporte
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(200, 750, "Reporte de Computadoras")

    # Encabezados
    pdf.setFont("Helvetica-Bold", 12)
    pdf.drawString(50, 720, "No. Inventario")
    pdf.drawString(200, 720, "Estado")
    pdf.drawString(350, 720, "Empleado")

    y = 700  # Posición inicial de los datos

    computadoras = Computadora.query.all()
    pdf.setFont("Helvetica", 10)

    for comp in computadoras:
        empleado_nombre = comp.empleado.nombre if comp.empleado else "No asignado"

        pdf.drawString(50, y, comp.numero_inventario)
        pdf.drawString(200, y, comp.estado)
        pdf.drawString(350, y, empleado_nombre)

        y -= 20  # Mover hacia abajo

        # Verifica si queda espacio en la página, si no, crea una nueva
        if y < 50:
            pdf.showPage()
            pdf.setFont("Helvetica", 10)
            y = 750  # Reiniciar posición en nueva página

    pdf.save()

    # Devolver el PDF como respuesta
    buffer.seek(0)
    return Response(buffer, mimetype='application/pdf', headers={"Content-Disposition": "attachment;filename=reporte_computadoras.pdf"})


from flask import session  # Asegúrate de importar session

@app.route('/reporte/excel')
def generar_reporte_excel():
    usuario_id = session.get('usuario_id')  # Obtener el usuario logueado
    usuario_rol = session.get('rol')  # Obtener el rol del usuario
    print("Usuario ID en sesión:", session.get('usuario_id'))
    print("Rol en sesión:", session.get('rol'))
    if not usuario_id:  # Si no hay usuario en sesión, no hacer nada
        return "No autorizado", 403

    # Si es admin, descargar todas las computadoras
    if usuario_rol == "admin":
        computadoras = Computadora.query.all()
    else:  
        # Si es empleado, descargar solo sus computadoras
        computadoras = Computadora.query.filter_by(empleado_id=usuario_id).all()

    # Convertir los datos a una lista de diccionarios
    data = [{
        "Número de Inventario": comp.numero_inventario,
        "Estado": comp.estado,
        "Departamento": comp.departamento,
        "Ubicación Campus": comp.ubicacion_campus,
        "Monitor": comp.monitor,
        "Teclado": comp.teclado,
        "Empleado": comp.empleado.nombre if comp.empleado else "No asignado"
    } for comp in computadoras]

    # Crear DataFrame y generar Excel
    df = pd.DataFrame(data)
    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name="Computadoras", index=False)

    # Devolver el archivo como respuesta
    excel_buffer.seek(0)
    return Response(
        excel_buffer,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment;filename=reporte_computadoras.xlsx"}
    )



# restringir el acceso a usuarios normales
@app.route('/obsoletas')
def pagina_obsoletas():
    if 'rol' not in session or session['rol'] != 'admin':
        return redirect(url_for('dashboard'))  # Redirige a usuarios normales

    return render_template('obsoletas.html')

@app.route('/computadoras/<int:id>/obsoleta', methods=['PUT'])
def marcar_como_obsoleta(id):
    computadora = Computadora.query.get(id)
    if not computadora:
        return jsonify({'error': 'Computadora no encontrada'}), 404

    computadora.es_obsoleta = True
    db.session.commit()

    return jsonify({'mensaje': 'Computadora marcada como obsoleta'}), 200
@app.route('/computadoras/obsoletas', methods=['GET'])
def obtener_obsoletass():
    computadoras = Computadora.query.filter_by(es_obsoleta=True).all()
    computadoras_json = [
        {
            'id': comp.id,
            'numero_inventario': comp.numero_inventario,
            'estado': comp.estado,
            'departamento': comp.departamento,
            'ubicacion_campus': comp.ubicacion_campus,
            'monitor': comp.monitor,
            'teclado': comp.teclado
        } for comp in computadoras
    ]
    return jsonify(computadoras_json)

@app.route('/reasignar', methods=['GET', 'POST'])
def reasignar_computadora():
    if request.method == 'POST':
        data = request.get_json()  # Obtener datos en formato JSON
        
        if not data:
            return jsonify({"success": False, "message": "No se enviaron datos"}), 400
        
        # Crear un nuevo objeto Reasignacion con los datos del formulario
        nueva_reasignacion = Reasignacion(
            tipo_equipo=data.get("tipo_equipo"),
            nomenclatura=data.get("nomenclatura"),
            service_tag=data.get("service_tag"),
            serie=data.get("serie"),
            marca=data.get("marca"),
            modelo=data.get("modelo"),
            inventario=data.get("inventario"),
            fecha_traslado=datetime.strptime(data.get("fecha_traslado"), '%Y-%m-%d'),
            propiedad=data.get("propiedad"),
            accesorios=data.get("accesorios"),
            observaciones=data.get("observaciones"),
            unidad_emisor=data.get("unidad_emisor"),
            responsable_emisor=data.get("responsable_emisor"),
            correo_emisor=data.get("correo_emisor"),
            empleado_emisor=data.get("empleado_emisor"),
            ubicacion_emisor=data.get("ubicacion_emisor"),
            unidad_receptor=data.get("unidad_receptor"),
            responsable_receptor=data.get("responsable_receptor"),
            correo_receptor=data.get("correo_receptor"),
            empleado_receptor=data.get("empleado_receptor"),
            ubicacion_receptor=data.get("ubicacion_receptor")
        )

        # Guardar en la base de datos
        db.session.add(nueva_reasignacion)
        db.session.commit()
        
        # Asegurar que el ID se está enviando en la respuesta
        return jsonify({
            "success": True,
            "message": "Reasignación guardada exitosamente.",
            "id": nueva_reasignacion.id  # Enviar el ID correcto
        })

    return render_template('reasignar.html')

@app.route('/generar_pdf/<int:id>', methods=['GET'])
def generar_pdf(id):
    # Obtener la reasignación desde la base de datos
    reasignacion = Reasignacion.query.get(id)

    if not reasignacion:
        return "Error: No se encontró la reasignación con ID {}".format(id), 404

    # Crear el PDF
    pdf_path = f"reasignacion_{id}.pdf"
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter

    # Encabezado
    c.setFont("Helvetica-Bold", 14)
    c.drawString(200, height - 50, "FORMATO DE REASIGNACIÓN DE COMPUTADORA")

    # Sección 1: Información del Equipo
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, height - 100, "I. Información del Equipo")
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 120, f"Tipo de Equipo: {reasignacion.tipo_equipo}")
    c.drawString(300, height - 120, f"Marca: {reasignacion.marca}")
    c.drawString(50, height - 140, f"Nomenclatura: {reasignacion.nomenclatura}")
    c.drawString(300, height - 140, f"Modelo: {reasignacion.modelo}")
    c.drawString(50, height - 160, f"Service Tag: {reasignacion.service_tag}")
    c.drawString(300, height - 160, f"Serie: {reasignacion.serie}")
    c.drawString(50, height - 180, f"Inventario: {reasignacion.inventario}")
    c.drawString(300, height - 180, f"Fecha de Traslado: {reasignacion.fecha_traslado.strftime('%Y-%m-%d')}")
    c.drawString(50, height - 200, f"Propiedad: {reasignacion.propiedad}")
    c.drawString(50, height - 220, f"Accesorios: {reasignacion.accesorios}")

    # Sección 2: Usuario Emisor
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, height - 260, "II. Información del Usuario Emisor")
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 280, f"Unidad/Depto: {reasignacion.unidad_emisor}")
    c.drawString(300, height - 280, f"Responsable: {reasignacion.responsable_emisor}")
    c.drawString(50, height - 300, f"Correo: {reasignacion.correo_emisor}")
    c.drawString(300, height - 300, f"Número de Empleado: {reasignacion.empleado_emisor}")
    c.drawString(50, height - 320, f"Ubicación Física: {reasignacion.ubicacion_emisor}")

    # Sección 3: Usuario Receptor
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, height - 360, "III. Información del Usuario Receptor")
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 380, f"Unidad/Depto: {reasignacion.unidad_receptor}")
    c.drawString(300, height - 380, f"Responsable: {reasignacion.responsable_receptor}")
    c.drawString(50, height - 400, f"Correo: {reasignacion.correo_receptor}")
    c.drawString(300, height - 400, f"Número de Empleado: {reasignacion.empleado_receptor}")
    c.drawString(50, height - 420, f"Ubicación Física: {reasignacion.ubicacion_receptor}")

    # Sección 4: Observaciones
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, height - 460, "IV. Observaciones")
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 480, f"{reasignacion.observaciones}")

    # Espacio para firmas
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, height - 520, "_________________________")
    c.drawString(300, height - 520, "_________________________")
    c.setFont("Helvetica", 10)
    c.drawString(50, height - 535, "Firma del Usuario Emisor")
    c.drawString(300, height - 535, "Firma del Usuario Receptor")

    # Guardar el PDF
    c.save()

    # Enviar el archivo al usuario
    return send_file(pdf_path, as_attachment=True)



@app.route('/generar_word/<int:id>')
def generar_word(id):
    # Obtener la reasignación desde la base de datos
    reasignacion = Reasignacion.query.get(id)

    if not reasignacion:
        return "Error: No se encontró la reasignación con ID {}".format(id), 404

    # Crear documento de Word
    doc = Document()

    # Sección 1: Información del Equipo
    p = doc.add_paragraph()
    run = p.add_run("I. INFORMACIÓN DEL EQUIPO")
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = 1  # Centrar texto

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    # Ajustar anchos de columna
    col_widths = [2000, 3000, 2000, 3000]

    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Rellenar la tabla
    data = [
        ("Tipo de Equipo:", reasignacion.tipo_equipo, "Marca:", reasignacion.marca),
        ("Nomenclatura:", reasignacion.nomenclatura, "Modelo:", reasignacion.modelo),
        ("Service TAG:", reasignacion.service_tag, "Inventario:", reasignacion.inventario),
        ("Serie:", reasignacion.serie, "Fecha de Traslado:", reasignacion.fecha_traslado.strftime('%Y-%m-%d')),
        ("Condiciones del equipo:", "Nuevo (  )  Usado (  )", "Propiedad:", "Leasing (  )  Propio (  )"),
    ]


    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Eliminar la fila extra y corregir "Componentes o Accesorios Entregados" y "Observaciones"
    table.add_row()
    table.cell(5, 0).merge(table.cell(5, 1))  # Convertir en 2 columnas
    table.cell(5, 2).merge(table.cell(5, 3))  # Convertir en 2 columnas
    p = table.cell(5, 0).paragraphs[0]
    p.add_run("Componentes o Accesorios Entregados:").bold = True
    table.cell(5, 2).text = reasignacion.accesorios if reasignacion.accesorios else "____________________"

    table.add_row()
    table.cell(6, 0).merge(table.cell(6, 1))  # Convertir en 2 columnas
    table.cell(6, 2).merge(table.cell(6, 3))  # Convertir en 2 columnas
    p = table.cell(6, 0).paragraphs[0]
    p.add_run("Observaciones:").bold = True
    table.cell(6, 2).text = reasignacion.observaciones if reasignacion.observaciones else "____________________"

    # Ajustar la altura de las filas para que sean más grandes
    table.rows[5].height = Pt(50)  # Hacer más grande "Componentes o Accesorios"
    table.rows[6].height = Pt(50)  # Hacer más grande "Observaciones"
    # 🔥 ELIMINAR FILAS EXCEDENTES 🔥
    while len(table.rows) > 7:  # Asegurar que solo haya 7 filas en total
     table._element.remove(table.rows[-1]._element)

# 🔥 NO AGREGAMOS MÁS FILAS DESPUÉS DE OBSERVACIONES 🔥
 
    # Sección 2: Información Usuario Emisor
    p = doc.add_paragraph()
    run = p.add_run("II. INFORMACIÓN USUARIO EMISOR")
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = 1  # Centrar texto

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    data = [
        ("Unidad o Depto.:", reasignacion.unidad_emisor),
        ("Responsable que entrega el equipo:", reasignacion.responsable_emisor),
        ("Correo Institucional:", f"{reasignacion.correo_emisor}    No. Empleado: {reasignacion.empleado_emisor}"),
        ("Ubicación Física:", reasignacion.ubicacion_emisor)
    ]

    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Sección 3: Información Usuario Receptor
    p = doc.add_paragraph()
    run = p.add_run("III. INFORMACIÓN USUARIO RECEPTOR")
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = 1  # Centrar texto

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    data = [
        ("Unidad o Depto.:", reasignacion.unidad_receptor),
        ("Responsable que recibe el equipo:", reasignacion.responsable_receptor),
        ("Correo Institucional:", f"{reasignacion.correo_receptor}    No. Empleado: {reasignacion.empleado_receptor}"),
        ("Ubicación Física:", reasignacion.ubicacion_receptor)
    ]

    for row_idx, row_data in enumerate(data):
        for col_idx, text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Agregar el código "GT FRM-001" fuera de la tabla
    p = doc.add_paragraph("GT FRM-001")
    p.alignment = 2  # Alineado a la derecha

    # Sección de Firmas
    p = doc.add_paragraph()
    run = p.add_run("USUARIO EMISOR:")
    run.bold = True
    run.font.size = Pt(12)

    run = p.add_run("\t\t\t\t\t\t")  # Espaciado para alinear
    run = p.add_run("USUARIO RECEPTOR:")
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.add_run("Nombre: __________________________\t\t\t\t\tNombre: __________________________").font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run("No. Empleado: ____________________\t\t\t\t\tNo. Empleado: ____________________").font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run("Firma: ___________________________\t\t\t\t\tFirma: ___________________________").font.size = Pt(10)

    # Guardar el documento en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name=f"reasignacion_{id}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
# Asegurar que la aplicación se ejecute correctamente

@app.route('/marcar_componente_obsoleto', methods=['POST'])
def marcar_componente_obsoleto():
    computadora_id = request.form.get('computadora_id')
    tipo_componente = request.form.get('tipo_componente')  # 'monitor' o 'teclado'
    motivo = request.form.get('motivo')  # Razón de obsolescencia

    computadora = Computadora.query.get(computadora_id)
    
    if not computadora:
        return jsonify({'message': 'Computadora no encontrada'}), 404

    # Determinar qué componente estamos marcando como obsoleto
    if tipo_componente == 'monitor':
        modelo_componente = computadora.monitor
        computadora.monitor = None  # Se vacía el campo en la tabla `computadoras`
        computadora.monitor_obsoleto = True  # Se marca como obsoleto
    elif tipo_componente == 'teclado':
        modelo_componente = computadora.teclado
        computadora.teclado = None  # Se vacía el campo en la tabla `computadoras`
        computadora.teclado_obsoleto = True  # Se marca como obsoleto
    else:
        return jsonify({'message': 'Tipo de componente inválido'}), 400

    # Guardar el componente en la tabla `componentes_defectuosos`
    componente_defectuoso = ComponenteDefectuoso(
        tipo_componente=tipo_componente.capitalize(),
        modelo=modelo_componente,
        motivo=motivo,
        computadora_id=computadora.id
    )
    db.session.add(componente_defectuoso)

    # Guardar cambios en la base de datos
    db.session.commit()

    return jsonify({'message': f'{tipo_componente.capitalize()} marcado como obsoleto correctamente'}), 200


#---ruta para empleados----
@app.route('/empleados_page')
@login_required
def empleados_page():
    return render_template('empleados.html', rol=current_user.rol)

#---ruta de accesorios----
@app.route('/accesorios_generales')
@login_required
def vista_accesorios_generales():
    if current_user.rol != 'admin':
        return redirect('/dashboard')
    
    # Versión fragmentada para carga dinámica
    if request.args.get('fragment'):
        return render_template('accesorios_generales.html')
    
    # Versión completa para carga normal
    return render_template('accesorios_generales.html')

@app.route('/empleados_json')
@login_required
def obtener_empleados_json():
    empleados = Empleado.query.all()
    return jsonify([
        {
            'id': emp.id,
            'nombre': emp.nombre,
            'email': emp.email
        }
        for emp in empleados
    ])

@app.route('/accesorios_generales/<int:empleado_id>', methods=['GET'])
@login_required
def obtener_accesorios_generales(empleado_id):
    empleado = Empleado.query.get(empleado_id)
    if not empleado:
        return jsonify({'error': 'Empleado no encontrado'}), 404

    accesorios = Accesorio.query.filter_by(empleado_id=empleado_id, computadora_id=None).all()

    return jsonify({
        'empleado': {
            'nombre': empleado.nombre,
            'email': empleado.email,
            'campus': empleado.campus
        },
        'accesorios': [
            {
                'id': a.id,
                'tipo': a.tipo,
                'numero_inventario': a.numero_inventario
            } for a in accesorios
        ]
    })


@app.route('/accesorios_generales', methods=['POST'])
@login_required
def agregar_accesorio_general():
    data = request.get_json()
    accesorio = Accesorio(
        tipo=data['tipo'],
        numero_inventario=data['numero_inventario'],
        empleado_id=data['empleado_id'],
        computadora_id=None
        
    )
    db.session.add(accesorio)
    db.session.commit()
    return jsonify({'mensaje': 'Accesorio agregado correctamente'}), 201


@app.route('/accesorios_generales/<int:accesorio_id>', methods=['DELETE'])
@login_required
def eliminar_accesorio_general(accesorio_id):
    accesorio = Accesorio.query.get(accesorio_id)
    if not accesorio:
        return jsonify({'error': 'Accesorio no encontrado'}), 404

    data = request.get_json()
    motivo = data.get('motivo', '').strip()

    if not motivo:
        return jsonify({'error': 'Motivo requerido'}), 400

    componente_obsoleto = ComponenteDefectuoso(
        tipo_componente=accesorio.tipo,
        modelo=accesorio.tipo,
        inventario_componente=accesorio.numero_inventario,
        motivo=motivo,
        computadora_id=None,
        empleado_id=accesorio.empleado_id
    )
    db.session.add(componente_obsoleto)
    db.session.delete(accesorio)
    db.session.commit()

    return jsonify({'mensaje': 'Accesorio marcado como obsoleto'}), 200


@app.route('/base_demo')
@login_required
def ver_base_demo():
    return render_template('base.html')

@app.route("/conexion")
def probar_conexion():
    try:
        db.session.execute(text("SELECT 1"))  # 👈 Aquí usamos text()
        return "✅ Conexión exitosa con PostgreSQL"
    except Exception as e:
        return f"❌ Error: {e}"

@app.route('/reporte/componentes_obsoletos/excel')
@login_required
def exportar_componentes_obsoletos_excel():
    if current_user.rol != "admin":
        return "Acceso denegado", 403

    # Obtener filtros desde los parámetros GET
    inventario = request.args.get('inventario')
    tipo = request.args.get('tipo')
    fecha_desde = request.args.get('fecha_desde')

    # Consulta base
    query = ComponenteDefectuoso.query

    if inventario:
        query = query.filter(ComponenteDefectuoso.inventario_componente.ilike(f"%{inventario}%"))
    if tipo:
        query = query.filter_by(tipo_componente=tipo)
    if fecha_desde:
        try:
            fecha_desde = datetime.fromisoformat(fecha_desde)
            query = query.filter(ComponenteDefectuoso.fecha_marcado >= fecha_desde)
        except ValueError:
            pass

    componentes = query.order_by(ComponenteDefectuoso.fecha_marcado.desc()).all()

    # Crear DataFrame
    df = pd.DataFrame([{
    "Tipo": c.tipo_componente,
    "Inventario": c.inventario_componente,
    "Modelo": c.modelo,
    "Motivo": c.motivo,
    "Fecha": c.fecha_marcado.strftime('%Y-%m-%d %H:%M'),
    "Computadora": c.computadora.numero_inventario if c.computadora else "No encontrada",
    "Empleado": (
        c.computadora.empleado.nombre if c.computadora and c.computadora.empleado
        else c.empleado.nombre if c.empleado else "Sin asignar"
    )
}   for c in componentes])

    # Convertir a Excel
    from io import BytesIO
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Obsoletos')

    output.seek(0)

    from flask import send_file
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        download_name='reporte_obsoletos.xlsx',
        as_attachment=True
    )


@app.route('/reporte/exportar_excel_eliminadas')
@login_required
def exportar_excel_eliminadas():
    if current_user.rol != 'admin':
        return jsonify({'error': 'Acceso denegado'}), 403

    texto = request.args.get('texto', '').strip().lower()
    tipo = request.args.get('tipo', '').strip()
    estado = request.args.get('estado', '').strip()


    query = ComputadoraEliminada.query

    from sqlalchemy import or_

    if texto:
        query = query.filter(
        or_(
            ComputadoraEliminada.numero_inventario.ilike(f'%{texto}%'),
            ComputadoraEliminada.departamento.ilike(f'%{texto}%'),
            ComputadoraEliminada.ubicacion_campus.ilike(f'%{texto}%'),
            ComputadoraEliminada.empleado_nombre.ilike(f'%{texto}%')
        )
    )

    if tipo:
        query = query.filter_by(tipo_propiedad=tipo)
    if estado:
        query = query.filter_by(estado=estado)

    datos = query.all()

    df = pd.DataFrame([{
        "N° Inventario": c.numero_inventario,
        "Estado": c.estado,
        "Departamento": c.departamento,
        "Campus": c.ubicacion_campus,
        "Monitor": c.monitor,
        "Teclado": c.teclado,
        "Propiedad": c.tipo_propiedad,
        "Fecha Eliminación": c.fecha_eliminacion.strftime('%Y-%m-%d %H:%M'),
        "Empleado": c.empleado_nombre or "No asignado"
    } for c in datos])

    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Eliminadas')

    output.seek(0)
    return send_file(output, download_name="eliminadas.xlsx", as_attachment=True)


if __name__ == '__main__':
    from flask_migrate import upgrade
    with app.app_context():
        upgrade()



#if __name__ == '__main__':
 #   with app.app_context():
  #   db.create_all() 
  #   app.run(debug=True)